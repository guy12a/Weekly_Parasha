from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Pt
from docx.shared import Inches
import requests
import re

# Precompile for speed
# This covers [יהוה, אדני, אל, אלוה, אלוהים, אלהים, צבאות, שדי]
HEBREW_VOWELS = r"[\u0591-\u05C7]*"
BOUNDARY = r"(?=$|[\s.,;:!?()\[\]{}׳״\-])"

DIVINE_NAMES_PATTERN = re.compile(
    (
        # 1. יהוה
        r"י" + HEBREW_VOWELS +
        r"ה" + HEBREW_VOWELS +
        r"ו" + HEBREW_VOWELS +
        r"ה" + HEBREW_VOWELS +

        # 2. אלוהים / אלהים
        r"|א" + HEBREW_VOWELS +
        r"ל" + HEBREW_VOWELS +
        r"ה" + HEBREW_VOWELS +
        r"י" + HEBREW_VOWELS +
        r"ם" + HEBREW_VOWELS +

        # 3. אלוה
        r"|א" + HEBREW_VOWELS +
        r"ל" + HEBREW_VOWELS +
        r"ו" + HEBREW_VOWELS +
        r"ה" + HEBREW_VOWELS +

        # 4. אדני
        r"|א" + HEBREW_VOWELS +
        r"ד" + HEBREW_VOWELS +
        r"נ" + HEBREW_VOWELS +
        r"י" + HEBREW_VOWELS +

        # 5. צבאות
        r"|צ" + HEBREW_VOWELS +
        r"ב" + HEBREW_VOWELS +
        r"א" + HEBREW_VOWELS +
        r"ו" + HEBREW_VOWELS +
        r"ת" + HEBREW_VOWELS +

        # 6. שדי
        r"|ש" + HEBREW_VOWELS +
        r"ד" + HEBREW_VOWELS +
        r"י" + HEBREW_VOWELS +

        # 7. Standalone אל (with boundary)
        r"|א" + HEBREW_VOWELS +
        r"ל" + HEBREW_VOWELS +
        BOUNDARY
    )
)


def replace_divine_names(text):
    print("Replacing names")
    return DIVINE_NAMES_PATTERN.sub("אלוקים", text)

# Takes the parasha range (genesis 1:1-10:2 for example), the language
def format(parasha_ref, lang):
    # Extracts the first numnber of the verse
    match = re.search(r":(\d+)", parasha_ref)
    if match:
        counter = int(match.group(1))

    # Get the text of the relevant language
    result = []
    if(lang == "heb"):
        nested_text = get_hebrew(parasha_ref)
    else:
        nested_text = get_english(parasha_ref)

    # Edits the text to add numberings
    for element in nested_text:
        for line in element:
            result.append(f"({counter}) {line}")
            counter += 1
        if(lang=="heb"):
            result.append("{פ}")  # Add {end} after each element
        counter = 1

    # Join everything with a space between lines
    toReturn =  " ".join(result)

    # Replace the explicit name
    if(lang == "heb"):
        toReturn = replace_divine_names(toReturn)
    else:
        toReturn = toReturn.replace("יהוה", "HaShem")

    return toReturn

# Make the hebreww text right leaning (somehow)
def make_run_rtl(run):
    """Mark a run (text) as RTL."""
    r = run._element
    rPr = r.get_or_add_rPr()
    
    # Add RTL property
    rtl = OxmlElement('w:rtl')
    rtl.set(qn('w:val'), '1')
    rPr.append(rtl)
    
    # Add font size at XML level
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '30')  # 32 half-points = 16 points
    rPr.append(sz)
    
    szCs = OxmlElement('w:szCs')  # Complex script size (for RTL text)
    szCs.set(qn('w:val'), '30')
    rPr.append(szCs)

# Get the hebrew version of the text
def get_hebrew(parasha):
    print("Getting the Hebrew version")
    hebURL = f"https://www.sefaria.org/api/v3/texts/{parasha}?version=hebrew%7CTanach%20with%20Nikkud&return_format=text_only"
    hebDATA = (requests.get(hebURL)).json()
    he_vtitle = hebDATA['versions'][0]['versionTitle'] 
    he_pasuk = hebDATA['versions'][0]['text']
    return he_pasuk

# Get the english version of the text
def get_english(parasha):
    print("Getting the English version")
    engURL = f"https://www.sefaria.org/api/v3/texts/{parasha}?version=english%7CThe%20Contemporary%20Torah%2C%20Jewish%20Publication%20Society%2C%202006&return_format=text_only"
    engDATA = (requests.get(engURL)).json()
    en_vtitle = engDATA['versions'][0]['versionTitle'] 
    en_pasuk = engDATA['versions'][0]['text']
    # print(en_pasuk) # The name of the Parasha
    return en_pasuk

# Create a document, fill it with text, and save it 
def construct_document_base(parasha_name,parasha_ref ,parasha_ref_heb):
    print("Creating the document")

    doc = Document()

    title = doc.add_paragraph()
    titlerun = title.add_run(parasha_name)
    titlerun.bold = True
    titlerun.underline=True
    titlerun.font.size = Pt(24)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    table = doc.add_table(rows=1, cols=2)

    table.autofit = False
    table.allow_autofit = False

    # Set column widths
    table.columns[0].width = Inches(3)   # half of the page width
    table.columns[1].width = Inches(3)

    row = table.rows[0]

    # LEFT COLUMN (English)
    left_cell = row.cells[0]
    left_cell._element.clear_content()

    title_eng = left_cell.add_paragraph()
    title_run_eng = title_eng.add_run(parasha_ref)
    title_run_eng.bold = True
    title_run_eng.font.size = Pt(17)
    title_eng.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    p_left = left_cell.add_paragraph()
    p_left.text = format(parasha_ref, "eng")
    p_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # p_left = left_cell.paragraphs[0]
    # p_left.text = format(parasha_ref,"eng")
    # p_left.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    for run in p_left.runs:
        run.font.size = Pt(13) 

    # RIGHT COLUMN (Hebrew)
    right_cell = row.cells[1]

    right_cell._element.clear_content()

    title_heb = right_cell.add_paragraph()
    title_run_heb = title_heb.add_run(parasha_ref_heb)
    title_run_heb.bold = True
    title_run_heb.font.size = Pt(17)
    title_heb.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT


    p_right = right_cell.add_paragraph()
    p_right.text = format(parasha_ref,"heb")
    # p_right.text = hebrewText.replace("יְהוָה", "HaShem")
    p_right.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    #p_right = right_cell.paragraphs[0]
    #p_right.text = re.sub(r"[A-Za-z]", "", hebrewFormat(parasha_ref).replace("יהוה", "אלוקים"))
    for run in p_right.runs:
        make_run_rtl(run)

    doc.save(f"{parasha_name}.docx")

url = "https://www.sefaria.org/api/calendars"
data = (requests.get(url)).json()
calendar_items = data['calendar_items']
for item in calendar_items:
    if item['title']['en'] == 'Parashat Hashavua':
        parasha_ref = item['ref'] 
        parasha_ref_heb = item['heRef']
        parasha_name = item['displayValue']['en']
        parasha_name_heb = item['displayValue']

print(f"Creating this weeks parasha: {parasha_name}")
construct_document_base(parasha_name,parasha_ref,parasha_ref_heb)


# Repository of functions that I have better versions of now ^^
# ====================================================
# def hebrewFormat(parasha_ref):
#     # Swap inner brackets
#     def swap_brackets(s):
#         return s.replace('[', 'TEMP').replace(']', '[').replace('TEMP', ']')
    
#     text = get_hebrew(parasha_ref)
#     formatted_blocks = []
#     for block in text:
#         # Swap brackets inside each line
#         swapped_lines = [swap_brackets(line) for line in block]
#         # Join lines with a space
#         formatted_blocks.append(' '.join(swapped_lines))

#     # Join blocks with {פ} between them
#     return ' }פ{ '.join(formatted_blocks)

# ====================================================

# def englishFormat(nested_list):
#     formatted_chunks = []
    
#     for inner in nested_list:
#         # Join all lines in the inner list with a space
#         chunk = " ".join(inner)
#         formatted_chunks.append(chunk)
    
#     # Join all chunks with "{end}"
#     result = " {end} ".join(formatted_chunks)
#     return result
# ====================================================

# def make_rtl(paragraph):
#     """
#     Mark a paragraph as RTL for Hebrew/Arabic text.
#     """
#     p = paragraph._element
#     pPr = p.get_or_add_pPr()
#     bidi = OxmlElement('w:bidi')
#     bidi.set(qn('w:val'), '1')
#     pPr.append(bidi)
# ====================================================

#def hebrewFormat2(parasha_ref):
#     nested_text = get_hebrew(parasha_ref)
#     result = []
#     for element in nested_text:
#         counter = 1
#         for line in element:
#             line.replace("יהוה", "אלוקים")
#             result.append(f"({counter}) {line}")
#             counter += 1
#         result.append("{פ}")  # Add {end} after each element
#     # Join everything with a space between lines
#     return " ".join(result)
# ====================================================

# def englishFormat2(parasha_ref):
#     nested_text = get_english(parasha_ref)
#     result = []
#     for element in nested_text:
#         counter = 1
#         for line in element:
#             line.replace("יהוה", "HaShem")
#             result.append(f"({counter}) {line}")
#             counter += 1
#         result.append("{end}")  # Add {end} after each element
#     # Join everything with a space between lines
#     return " ".join(result)
# ====================================================

# def replaceTet(text):
#     text = text.replace("יְהוָה", "אלוקים")
#     text = text.replace("יהוָה", "אלוקים")
#     text = text.replace("יהֹוָה", "אלוקים")
#     text = text.replace("יְהֹוָה", "אלוקים")
#     text = text.replace("יהוה", "אלוקים")
#     return text